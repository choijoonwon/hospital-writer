import os
import io
from pathlib import Path
from datetime import datetime
from fastapi import FastAPI, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from typing import List, Optional
from dotenv import load_dotenv

BASE_DIR = Path(__file__).parent.parent
load_dotenv(BASE_DIR / ".env")

app = FastAPI(title="병원 마케팅 글쓰기 자동화")

PLATFORMS = ["여우야", "여우야+", "성위키", "성예사", "재잘재잘", "가아사", "웨딩킹", "바비", "강남언니"]
POST_TYPES = ["게시글", "댓글", "쪽지"]


class GenerateRequest(BaseModel):
    patient_name: str
    platform: str
    post_type: str
    extra_context: str = ""
    char_limit: int = 0
    hospital_format: str = "full"
    use_period: bool = True  # False면 문장 끝 온점 사용 안 함


@app.get("/api/patients")
def list_patients():
    try:
        from backend.sheets import get_patients
        patients = get_patients()
        return {"patients": [
            {
                "이름": p["이름"],
                "사용불가": p.get("사용불가", False),
                "병원명": p.get("병원명", ""),
                "수술부위": p.get("수술부위", ""),
                "수술날": p.get("수술날", ""),
                "원장님": p.get("원장님", ""),
                "말투": p.get("말투", ""),
                "리스트": p.get("리스트", ""),
                "특이사항": p.get("특이사항", ""),
                "비고": p.get("비고", ""),
                "년생": p.get("년생", ""),
                "추천인": p.get("추천인", ""),
                "카페목록": [c.get("닉네임", "") + " " + c.get("카페", "") for c in p.get("카페목록", [])],
            }
            for p in patients
        ]}
    except FileNotFoundError:
        raise HTTPException(status_code=503, detail="credentials.json 파일이 없습니다.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/hospitals")
def list_hospitals():
    try:
        from backend.sheets import get_patients
        patients = get_patients()
        hospitals = sorted({p.get("병원명", "") for p in patients if p.get("병원명")})
        return {"hospitals": hospitals}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/patients/{name}")
def get_patient(name: str):
    try:
        from backend.sheets import get_patient_by_name
        patient = get_patient_by_name(name)
        if not patient:
            raise HTTPException(status_code=404, detail="환자를 찾을 수 없습니다.")
        return patient
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/generate")
def generate(req: GenerateRequest):
    if req.post_type not in POST_TYPES:
        raise HTTPException(status_code=400, detail=f"지원하지 않는 글 종류: {req.post_type}")

    try:
        from backend.sheets import get_patient_by_name
        patient = get_patient_by_name(req.patient_name)
        if not patient:
            raise HTTPException(status_code=404, detail="환자를 찾을 수 없습니다.")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"시트 읽기 오류: {e}")

    try:
        from backend.generator import generate_post, review_post
        text = generate_post(patient, req.platform, req.post_type, req.extra_context, req.char_limit, req.hospital_format, req.use_period)
        review = review_post(text, patient, req.platform, req.post_type)
        return {"result": text, "review": review}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI 생성 오류: {e}")


@app.post("/api/patients/refresh")
def refresh_patients():
    try:
        from backend.sheets import refresh_cache
        patients = refresh_cache()
        return {"count": len(patients), "message": "새로고침 완료"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class ExportItem(BaseModel):
    name: str
    hospital: str = ""
    platform: str = ""
    post_type: str = ""
    text: str
    char_count: int = 0
    review_pass: Optional[str] = None   # "pass" | "warn" | "fail" | None


class ExportRequest(BaseModel):
    items: List[ExportItem]
    filename: str = "생성결과"


@app.post("/api/export-docx")
def export_docx(req: ExportRequest):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 여백 설정
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)

        # 제목
        today = datetime.now().strftime("%Y-%m-%d")
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"병원 마케팅 글쓰기 생성 결과")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"생성일: {today}  |  총 {len(req.items)}건")
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # 공백

        REVIEW_LABEL = {"pass": "✓ 통과", "warn": "△ 주의", "fail": "✗ 수정필요"}
        REVIEW_COLOR = {
            "pass": RGBColor(0x2d, 0x7a, 0x3a),
            "warn": RGBColor(0xb4, 0x53, 0x09),
            "fail": RGBColor(0xc6, 0x28, 0x28),
        }

        for i, item in enumerate(req.items):
            # 구분선
            if i > 0:
                doc.add_paragraph("─" * 50)

            # 헤더 (이름 + 병원 + 플랫폼)
            header_para = doc.add_paragraph()
            name_run = header_para.add_run(f"{item.name}")
            name_run.bold = True
            name_run.font.size = Pt(13)

            if item.hospital:
                hosp_run = header_para.add_run(f"  |  {item.hospital}")
                hosp_run.font.size = Pt(11)
                hosp_run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

            if item.platform:
                plat_run = header_para.add_run(f"  ·  {item.platform}")
                plat_run.font.size = Pt(10)
                plat_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            # 검수 결과 + 글자수
            meta_para = doc.add_paragraph()
            cnt_run = meta_para.add_run(f"{item.char_count}자")
            cnt_run.font.size = Pt(9)
            cnt_run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

            if item.review_pass and item.review_pass in REVIEW_LABEL:
                rev_run = meta_para.add_run(f"  {REVIEW_LABEL[item.review_pass]}")
                rev_run.bold = True
                rev_run.font.size = Pt(9)
                rev_run.font.color.rgb = REVIEW_COLOR[item.review_pass]

            # 본문
            body_para = doc.add_paragraph()
            body_run = body_para.add_run(item.text)
            body_run.font.size = Pt(11)

            doc.add_paragraph()  # 여백

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        from urllib.parse import quote
        safe_filename = quote(req.filename.encode("utf-8"))

        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}.docx"},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX 생성 오류: {e}")


@app.get("/api/config")
def get_config():
    return {
        "platforms": PLATFORMS,
        "post_types": POST_TYPES,
        "ai_provider": os.getenv("AI_PROVIDER", "claude"),
        "sheets_configured": bool(os.getenv("SPREADSHEET_ID")),
    }


app.mount("/", StaticFiles(directory=str(BASE_DIR / "frontend"), html=True), name="frontend")
